# -*- coding: utf-8 -*-
"""Generate NestEdu architecture doc mimicking NestEdu_Go.docx structure."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT = Path(r"C:\Users\13207\Desktop\NestEdu_启芽智教.docx")


def setup_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(11)


def add_code_block(doc: Document, text: str) -> None:
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(10)


def build() -> Document:
    doc = Document()
    setup_styles(doc)

    # Title
    t = doc.add_paragraph()
    tr = t.add_run("NestEdu 华科附幼智能教案助手（启芽智教）— 架构与核心技术实现")
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.name = "微软雅黑"
    tr._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    add_heading(doc, "目录", 1)
    toc = [
        "1 项目整体介绍",
        "  1.1 项目定位与简历总述",
        "  1.2 工程多模块结构",
        "  1.3 整体架构",
        "2 四大核心技术实现（原文完整保留）",
        "  2.1 Go BFF 分层架构",
        "  2.2 平台网关与鉴权透传",
        "  2.3 智能体调用与文档处理",
        "  2.4 工程化交付（pnpm monorepo + Docker 多阶段）",
        "3 高频面试问答",
        "  3.1 这个项目后端 Go 具体干什么？",
        "  3.2 为什么叫 BFF？和普通后端有什么区别？",
        "  3.3 你们怎么登录、Token 怎么传到平台？",
        "  3.4 AI 生成是在你服务里跑还是调平台？",
        "  3.5 知识库和本地 MySQL 分别存什么？",
        "  3.6 生产环境怎么部署？本地为什么可以不启 Go？",
    ]
    for line in toc:
        add_para(doc, line)

    # === 1 项目整体介绍 ===
    add_heading(doc, "1 项目整体介绍", 1)

    add_heading(doc, "1.1 项目定位与简历总述", 2)
    add_para(
        doc,
        "NestEdu（启芽智教）面向华中科技大学附属幼儿园一线教师的智能工作与成长平台，"
        "在 AI101 智能体与知识库之上，把活动方案生成、周计划编排、成果沉淀、教师画像收成同一套鉴权与数据闭环；"
        "基于 Go+Gin BFF + AI101 智能体对接 + 知识库三分类管理 + 双通道鉴权（iframe / SSO）+ 同域反代与单镜像交付。",
    )
    add_para(doc, "简历总述原文：", bold=True)
    add_para(
        doc,
        "NestEdu（华科附幼智能教案助手 / 启芽智教）面向华科附幼的教学助手，"
        "基于主题自动生成活动方案与周计划，并支持成果库多格式沉淀与教师画像解读；"
        "基于 Go+Gin 搭建 BFF，对接 AI101 智能体与知识库，完成方案生成、周计划编排、文档入库与 Word/PDF 导出；"
        "生产环境由 Go 统一托管静态资源并反向代理平台 API。",
    )

    add_heading(doc, "1.2 工程多模块结构", 2)
    add_code_block(
        doc,
        """
NestEdu (pnpm monorepo)
├── apps/web/           # React 19 + TypeScript + Vite — 页面、hooks、api、lib
├── apps/api/           # Go + Gin BFF — http / service / store / model
├── docker/             # 多阶段镜像
├── docs/               # 架构、接口、认证说明
├── openspec/           # 变更规格
├── scripts/            # CI、OpenSpec 脚本
├── AGENTS.md / CHANGELOG.md
└── .env.example
        """,
    )

    add_heading(doc, "1.3 整体架构", 2)
    add_para(
        doc,
        "整体是「浏览器 — NestEdu 前端 — Go BFF（可选）— 平台 AI101」的教学工作台："
        "前端按 pages / hooks / api / lib 分层；后端按 http / service / store / model 分层。"
        "平台侧负责智能体推理与知识库存储；本地 MySQL 可选持久化本人入库计数、成果录入与画像快照。",
    )

    add_para(doc, "1.3.1 鉴权与请求链路", bold=True)
    add_para(
        doc,
        "用户登录后统一走 auth-bridge：嵌入 AI101 iframe 时向父窗口取 Token；顶层直访时走 SSO（ticket 换 Token）。"
        "axios 拦截器自动注入 Authorization、X-Uid-Hash 等头。生产同域走 Go：/api/v1 走本地 BFF，"
        "/api/knowledge、/api/file、/api/user、/v1 由 Go 反代 api.zcat.cn；开发期 Vite 代理直连平台，可不启 Go。",
    )

    add_para(doc, "1.3.2 活动方案生成主链路", bold=True)
    add_para(
        doc,
        "用户选班级、重点领域、主题 → 前端拼装提示词 → 调平台 POST /v1/text/generate（教案 Agent 14317）"
        "→ 解析为结构化 TeachingPlan → 确认后写入知识库教案分类（20806）+ 可选 MySQL 本人计数。"
        "入库标题约定「姓名_活动方案_主题」，禁止误入教师成果库手机号文件夹。",
    )

    add_para(doc, "1.3.3 周计划编排与导出", bold=True)
    add_para(
        doc,
        "用户选班级/主题/周次，从教案库勾选素材 → 周计划 Agent（14332）生成「快乐一周」结构"
        "（自主学习/自主游戏/自主生活/自主运动四栏）→ 单元格编辑、AI 改稿 → 导出 Word/PDF，"
        "并写入周计划知识库分类（20807）。",
    )

    add_para(doc, "1.3.4 成果库与教师画像", bold=True)
    add_para(
        doc,
        "成果库对接知识库 10298「教师成果库」（20895），按登录手机号匹配同名个人文件夹，"
        "支持 Word/PDF/PPT/Excel/图片等多格式上传（先 /api/file/upload，再登记知识库）。"
        "教师画像聚合本人活动/周计划计数与成果库文档，调用画像 Agent（14372）生成解读，快照可落库 profile_snapshots。",
    )

    add_para(doc, "1.3.5 工程化与交付", bold=True)
    add_para(
        doc,
        "鉴权依赖 AI101；正式教案/周计划以平台知识库为准；本地 GORM+MySQL 做可选持久化；"
        "pnpm 统一脚本；Docker 多阶段单镜像交付；失败显式报错，禁止默认 Mock 成功。",
    )

    # === 2 四大核心技术 ===
    add_heading(doc, "2 四大核心技术实现", 1)

    add_heading(doc, "2.1 Go BFF 分层架构", 2)
    add_para(doc, "原文（简历表述，已按项目实际修订）：", bold=True)
    add_para(
        doc,
        "1. Go BFF 分层架构：基于 Gin 按 http / service / store / model 分层，"
        "对外提供周计划 CRUD、教师成果录入（growth_records）、本人入库计数（teacher_generated_docs）、"
        "画像快照（profile_snapshots）及知识库封装等 /api/v1 接口；"
        "KnowledgeService 在上传时强制纠正业务分类，避免误入教师成果库；"
        "GORM + MySQL 持久化业务数据（本地 Docker / 生产 RDS），协议层与业务逻辑解耦。",
    )
    add_para(
        doc,
        "NestEdu 的后端是 Go + Gin 的 BFF，服务于前端页面，而不是替代全量中台。"
        "http 只处理协议；service 做业务编排与平台适配；store 用 GORM 访问数据库；model 只放数据结构。"
        "router 启动时按 Store → Service → Handler 组装，对外统一 /api/v1，同时注册平台反代与静态资源托管。",
    )

    add_heading(doc, "2.2 平台网关与鉴权透传", 2)
    add_para(
        doc,
        "2. 平台网关与鉴权透传：生产环境由 Go 反向代理知识库、文件服务与智能体 API，同域收敛跨域；"
        "service 层适配平台接口并透传 Token；auth-bridge 支持 iframe 取 Token 与 SSO ticket 换票；"
        "axios 拦截器自动注入鉴权头，未登录拦截。",
    )
    add_para(
        doc,
        "platform_proxy.go 用 httputil.ReverseProxy 将 /api/knowledge、/api/file、/api/user、/v1 转到 PLATFORM_API_BASE_URL。"
        "lib/authBridge.ts 封装 @zcat-open/auth-bridge，父窗口 origin 白名单配置于 VITE_AI101_PARENT_ORIGINS。"
        "api/client.ts 每次请求 buildAuthHeaders，并回退 X-Uid-Hash 缓存，满足知识库分类接口要求。",
    )
    add_para(
        doc,
        "主链路一句话：登录拿 Token（iframe 或 SSO）→ axios 自动带 Header → 同域走 Go（BFF 或反代）→ Token 透传平台 → 业务成功或显式失败。",
    )

    add_heading(doc, "2.3 智能体调用与文档处理", 2)
    add_para(
        doc,
        "3. 智能体调用与文档处理：对接平台文本生成 API，活动方案 / 周计划 / 教师画像分 Agent 调用；"
        "实现提示词拼装与结构化解析，知识库三分类读写，多格式文件上传与自主导出。",
    )
    add_para(
        doc,
        "生成接口：POST /v1/text/generate，携带用户 Token 与 agent_id（14317 / 14332 / 14372）。"
        "知识库 10298 三分类：教案 20806、周计划 20807、教师成果库 20895（下挂手机号个人文件夹）。",
    )
    add_para(
        doc,
        "文档处理：活动/周计划上传以 mammoth 解析 Word 后走 document/text；"
        "成果库多格式先 /api/file/upload 再登记知识库；导出走 export-doc / export-pdf。",
    )
    add_para(
        doc,
        "主链路一句话：选条件 → 调 Agent → 结构化解析 → 确认写入对应分类知识库 → 上传或导出 Word/PDF。",
    )

    add_heading(doc, "2.4 工程化交付", 2)
    add_para(doc, "原文（简历表述）：", bold=True)
    add_para(
        doc,
        "4. 工程化交付：pnpm monorepo 统一构建与脚本；Docker 多阶段构建（前端打包 + Go 精简二进制 + Alpine 运行时），"
        "单镜像完成静态托管、BFF 与平台反代。",
    )
    add_para(
        doc,
        "pnpm workspace 统一 dev / build / lint / ci；Docker 三阶段：web（Node+Vite）、api-builder（Go -trimpath）、runtime（Alpine）。"
        "VITE_* 为构建期注入，改知识库 ID、Agent ID 需重新 build。OpenSpec 管理中大型变更。",
    )

    # === 3 面试问答 ===
    add_heading(doc, "3 高频面试问答", 1)

    add_heading(doc, "3.1 这个项目后端 Go 具体干什么？", 2)
    add_para(
        doc,
        "简答：Go 主要做三件事——（1）BFF，分层提供 /api/v1 周计划、成果录入、本人入库计数、画像快照与知识库封装；"
        "（2）生产反代 /api/knowledge、/api/file、/v1 到 AI101，同域收敛；（3）托管前端静态资源。"
        "本地开发可不启 Go，前端经 Vite 代理直连平台；需要 MySQL 统计时启 Go。",
    )

    add_heading(doc, "3.2 为什么叫 BFF？和普通后端有什么区别？", 2)
    add_para(
        doc,
        "简答：BFF 是面向前端的边界层，不是通用业务中台。它封装 AI101 知识库与智能体协议、做同域反代、"
        "可选本地持久化，以及把多平台接口收敛成 /api/v1。重点是「为前端服务」+「分层接口」，而不是替代 AI101。",
    )

    add_heading(doc, "3.3 你们怎么登录、Token 怎么传到平台？", 2)
    add_para(
        doc,
        "简答：对接 AI101 账号，用 @zcat-open/auth-bridge。iframe 从父窗口取 Token，顶层直访走 SSO ticket 换票；"
        "axios 自动注入鉴权头；知识库/智能体请求时 Token 透传平台；父窗口 origin 白名单；未登录显式拦截，失败不 Mock。",
    )

    add_heading(doc, "3.4 AI 生成是在你服务里跑还是调平台？", 2)
    add_para(
        doc,
        "简答：模型在 AI101 侧。我们调 POST /v1/text/generate，分 Agent；前端负责提示词拼装和结构化解析，"
        "确认后写知识库。Go 可封装接口和反代，推理不在本地 Go 进程里。",
    )

    add_heading(doc, "3.5 知识库和本地 MySQL 分别存什么？", 2)
    add_para(
        doc,
        "简答：平台知识库存正式活动方案/周计划文档与教师成果库文件（三分类 + 手机号文件夹）；"
        "本地 MySQL 存 teacher_generated_docs（本人入库计数）、growth_records（教师录入成果）、"
        "profile_snapshots（画像快照）、weekly_plans（周计划 CRUD）。产品真相源以知识库为准。",
    )

    add_heading(doc, "3.6 生产环境怎么部署？本地为什么可以不启 Go？", 2)
    add_para(
        doc,
        "简答：Docker 多阶段单镜像——前端产物 + Go BFF + Alpine 运行，同域提供页面、BFF 与平台反代。"
        "本地 VITE_USE_BACKEND_API=false 时，Vite 把知识库、智能体、换票代理到平台，主链路可不启 Go；"
        "需要 MySQL 统计/画像快照时需启 Go + MySQL。",
    )

    # === 附录 ===
    add_heading(doc, "附录：简历段落串联", 1)
    add_para(
        doc,
        "附录串联：鉴权 → 第 2 点（auth-bridge + 透传）\n"
        "后端系统 → 第 1 点（五入口 + 第 1 点 /api/v1 分层）\n"
        "生成与文档 → 第 3 点（多 Agent + 三分类知识库 + 导出）\n"
        "工程交付 → 第 4 点（monorepo + Docker 单镜像）",
    )
    add_para(
        doc,
        "原则：技术栈强调全栈与 BFF，不要写项目没有的 Redis 集群；可写 MySQL 可选持久化与大模型在平台侧。",
    )
    add_para(
        doc,
        "NestEdu 适合用来讲「机制覆盖」：华科附幼教师从主题到活动方案、周计划、成果沉淀、画像成长的完整闭环；"
        "对接 AI101 而非自建模型；生成、入库、导出一体化。",
    )
    add_para(
        doc,
        "部署形态：前后端同仓；生产 Go 托管前端并反代平台；本地可只起 Vite 走代理，需要统计时再启 Go + MySQL。",
    )
    add_para(
        doc,
        "一句话总结：这是一个对接 AI101 的幼儿园教师智能工作台 MVP，"
        "我负责业务栈与 BFF、鉴权与反代、全栈落地；AI 推理在平台，我们管编排、入库与交付。",
    )

    return doc


if __name__ == "__main__":
    doc = build()
    doc.save(OUT)
    print(f"Wrote: {OUT}")
